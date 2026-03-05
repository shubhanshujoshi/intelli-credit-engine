import streamlit as st
import pandas as pd
import numpy as np
import joblib
import shap
import json
import matplotlib.pyplot as plt
import os
import re
import requests
import io
import PyPDF2
from docx import Document
from datetime import datetime, timedelta
from num2words import num2words
from google import genai

# ----------------------------------------------------------
# PAGE CONFIG
# ----------------------------------------------------------

st.set_page_config(layout="wide")
st.title("IntelliCredit-X | AI Credit Decision Engine")

# ----------------------------------------------------------
# LOAD API KEYS
# ----------------------------------------------------------

GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
GNEWS_API_KEY = os.getenv("GNEWS_API_KEY")

if GOOGLE_API_KEY:
    client = genai.Client(api_key=GOOGLE_API_KEY)
else:
    client = None

# ----------------------------------------------------------
# AI DATA INGESTOR (PILLAR 1) - PDF EXTRACTION
# ----------------------------------------------------------

def extract_text_from_pdf(uploaded_file):
    """Extracts raw text from an uploaded PDF file."""
    reader = PyPDF2.PdfReader(uploaded_file)
    text = ""
    for page in reader.pages:
        extracted = page.extract_text()
        if extracted:
            text += extracted + "\n"
    return text

def parse_financials_with_gemini(text):
    """Uses Gemini to extract financial data from messy text."""
    if not client or not text.strip():
        return None

    prompt = """
    You are an expert financial analyst. Read the following text extracted from a company document.
    Extract the following metrics if present. Return ONLY a valid JSON format with these exact keys, using purely numeric values (no commas, no currency symbols). If a value is missing or you cannot determine it, use 0.
    Keys: "Revenue", "EBITDA", "Debt".
    
    Text:
    """ + text[:15000] # Limit tokens to avoid overflow

    try:
        response = client.models.generate_content(
            model="models/gemini-2.5-flash",
            contents=prompt
        )
        
        # Clean up response to ensure it's valid JSON
        result_text = response.text.replace("```json", "").replace("```", "").strip()
        return json.loads(result_text)
    except Exception as e:
        return None

# ----------------------------------------------------------
# GEMINI SENTIMENT FUNCTION
# ----------------------------------------------------------

def get_news_sentiment_score(company_name, combined_news_text):

    if not client or not combined_news_text.strip():
        return 0.0

    prompt = f"""
    You are an expert corporate credit risk analyst.

    Based on the following recent news (last 60 days) about '{company_name}',
    provide a sentiment score between -1.0 and +1.0.

    Negative = high credit risk
    Positive = financially strong
    Return ONLY the number.

    News:
    "{combined_news_text}"
    """

    try:
        response = client.models.generate_content(
            model="models/gemini-2.5-flash",
            contents=prompt
        )

        score_text = response.text.strip()
        number = re.findall(r'-?\d+\.?\d*', score_text)[0]
        return float(number)

    except:
        return 0.0

# ----------------------------------------------------------
# FETCH LAST 60 DAYS NEWS
# ----------------------------------------------------------

def fetch_last_60_days_news(company_name):

    if not GNEWS_API_KEY:
        return []

    end_date = datetime.utcnow()
    start_date = end_date - timedelta(days=60)

    url = (
        f"https://gnews.io/api/v4/search?"
        f"q={company_name}&"
        f"from={start_date.strftime('%Y-%m-%d')}&"
        f"to={end_date.strftime('%Y-%m-%d')}&"
        f"lang=en&"
        f"max=10&"
        f"apikey={GNEWS_API_KEY}"
    )

    try:
        response = requests.get(url)
        return response.json().get("articles", [])
    except:
        return []

# ----------------------------------------------------------
# CAM GENERATOR (WORD DOC)
# ----------------------------------------------------------

def generate_cam_word(company_name, revenue, ebitda, debt, prob, sentiment, adjusted_loan, interest_rate, articles):
    """Generates a structured Credit Appraisal Memo based on the Five Cs."""
    doc = Document()
    doc.add_heading(f'Credit Appraisal Memo (CAM)', 0)
    doc.add_paragraph(f"Company: {company_name.upper()}")
    doc.add_paragraph(f"Date: {datetime.utcnow().strftime('%Y-%m-%d')}")
    
    doc.add_heading('Executive Summary', level=1)
    # Check if threshold is loaded in scope; fallback to 0.5 if not
    global threshold
    current_threshold = threshold if 'threshold' in globals() else 0.5
    decision = "APPROVED" if prob <= current_threshold else "REJECTED"
    
    doc.add_paragraph(f"Final Decision: {decision}\nRecommended Limit: ₹ {adjusted_loan:,.0f}\nProposed Interest Rate: {interest_rate:.2f}%\nProbability of Default: {prob:.2%}")

    doc.add_heading('1. Character (Management & Sentiment)', level=1)
    doc.add_paragraph(f"AI News Sentiment Score (Last 60 Days): {sentiment:.2f}")
    if articles:
        doc.add_paragraph("Recent Intelligence:")
        for a in articles[:3]: # Add top 3 articles
            doc.add_paragraph(f"- {a.get('source', {}).get('name', 'News')}: {a.get('title', '')}", style='List Bullet')

    doc.add_heading('2. Capacity (Financial Health)', level=1)
    doc.add_paragraph(f"Revenue: ₹ {revenue:,.2f}\nEBITDA: ₹ {ebitda:,.2f}\nDebt: ₹ {debt:,.2f}")
    
    doc.add_heading('3. Capital (Leverage)', level=1)
    leverage = debt / ebitda if ebitda > 0 else 0
    doc.add_paragraph(f"Calculated Debt/EBITDA Ratio: {leverage:.2f}x")

    doc.add_heading('4. Conditions (Sector & Macro)', level=1)
    doc.add_paragraph("Macro-economic risk and sector headwinds have been factored into the base PD model.")
    
    doc.add_heading('5. Collateral', level=1)
    doc.add_paragraph("To be assessed based on specific asset hypothecation and standard LTV ratios.")

    # Save to memory buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ----------------------------------------------------------
# LOAD MODEL FILES
# ----------------------------------------------------------

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

@st.cache_resource
def load_model():
    return joblib.load(os.path.join(BASE_DIR, "financial_model.pkl"))

@st.cache_resource
def load_threshold():
    with open(os.path.join(BASE_DIR, "config.json"), "r") as f:
        return json.load(f)["best_threshold"]

@st.cache_resource
def load_feature_names():
    return joblib.load(os.path.join(BASE_DIR, "feature_names.pkl"))

@st.cache_resource
def load_explainer():
    model = joblib.load(os.path.join(BASE_DIR, "financial_model.pkl"))
    return shap.TreeExplainer(model)

model = load_model()
threshold = load_threshold()
feature_names = load_feature_names()
explainer = load_explainer()

# ----------------------------------------------------------
# INPUT SECTION & DATA INGESTOR
# ----------------------------------------------------------

st.header("1. Document Ingestion (Pillar 1)")
st.info("Upload a financial statement (PDF) to auto-extract metrics using AI.")

uploaded_file = st.file_uploader("Upload Annual Report or Financial Statement (PDF)", type=["pdf"])

# Initialize session state for auto-filling so we don't overwrite manual edits unnecessarily
if 'auto_rev' not in st.session_state:
    st.session_state.update({'auto_rev': 5e8, 'auto_ebitda': 1e8, 'auto_debt': 2e8})

# Only parse if a new file is uploaded
if uploaded_file is not None:
    if st.session_state.get('uploaded_filename') != uploaded_file.name:
        with st.spinner("AI Agent extracting financial data from document..."):
            raw_text = extract_text_from_pdf(uploaded_file)
            extracted_data = parse_financials_with_gemini(raw_text)
            
            if extracted_data:
                st.success("✅ Financials successfully extracted via Gemini AI!")
                st.session_state.auto_rev = float(extracted_data.get("Revenue", st.session_state.auto_rev))
                st.session_state.auto_ebitda = float(extracted_data.get("EBITDA", st.session_state.auto_ebitda))
                st.session_state.auto_debt = float(extracted_data.get("Debt", st.session_state.auto_debt))
                with st.expander("View Extracted JSON Data"):
                    st.json(extracted_data) 
            else:
                st.error("Could not parse JSON from document. Using default or manual inputs.")
        
        # Mark this file as processed
        st.session_state['uploaded_filename'] = uploaded_file.name


st.header("2. Financial Verification & Primary Insights")

col1, col2 = st.columns(2)

with col1:
    revenue = st.number_input("Revenue", value=st.session_state.auto_rev)
    ebitda = st.number_input("EBITDA", value=st.session_state.auto_ebitda)
    debt = st.number_input("Debt", value=st.session_state.auto_debt)
    interest_cov = st.number_input("Interest Coverage", value=2.5)
    gst = st.slider("GSTR-2A/3B Variance (Indian Context)", 0.0, 0.5, 0.1)

with col2:
    litigation = st.number_input("Litigation Count", value=1)
    sector = st.slider("Sector Risk", 0.0, 1.0, 0.3)
    mgmt = st.slider("Management Quality", 1.0, 10.0, 7.0)
    capacity = st.slider("Capacity Utilization", 0.0, 1.0, 0.8)

st.subheader("External AI News Risk Analysis")
company_name = st.text_input("Company Name", value="tata power")

# ----------------------------------------------------------
# PREDICTION
# ----------------------------------------------------------

if st.button("🔍 Analyze Credit Risk"):

    articles = []
    sentiment_score = 0.0

    with st.spinner("Fetching last 60 days news and analyzing sentiment..."):
        articles = fetch_last_60_days_news(company_name)

        combined_news = ""
        for article in articles:
            combined_news += article.get("title", "") + ". "
            combined_news += article.get("description", "") + ". "

        sentiment_score = get_news_sentiment_score(company_name, combined_news)

    # ------------------------------------------------------
    # SHOW ARTICLES ANALYZED
    # ------------------------------------------------------

    st.subheader("News Intelligence Layer")

    article_count = len(articles)

    if article_count > 0:
        st.success(f"Articles Analyzed (Last 60 Days): {article_count}")

        for article in articles:
            title = article.get("title", "No Title")
            source = article.get("source", {}).get("name", "Unknown Source")
            url = article.get("url", "")
            st.markdown(f"• **{source}** – [{title}]({url})")
    else:
        st.warning("No recent news articles found for this company.")

    # ------------------------------------------------------
    # ML PREDICTION
    # ------------------------------------------------------

    input_data = np.array([[ 
        revenue, ebitda, debt, interest_cov,
        gst, litigation, sentiment_score, sector,
        mgmt, capacity
    ]])

    input_df = pd.DataFrame(input_data, columns=feature_names)

    prob = model.predict_proba(input_df)[0][1]
    prob = prob - (0.05 * sentiment_score)
    prob = max(0, min(prob, 1))

    prediction = int(prob > threshold)

    st.subheader("Risk Assessment")

    st.metric("Probability of Default (PD)", round(prob, 3))
    st.metric("AI News Sentiment Score (60D)", round(sentiment_score, 2))
    st.write("Decision Threshold:", round(threshold, 3))

    if prediction == 1:
        st.error("🚫 HIGH RISK — Loan Rejected")
    else:
        st.success("✅ LOW RISK — Loan Approved")

    # ------------------------------------------------------
    # LOAN PRICING LOGIC
    # ------------------------------------------------------

    max_loan = ebitda * 3.5
    adjusted_loan = max_loan * (1 - prob)

    base_rate = 9
    risk_premium = prob * 6
    interest_rate = base_rate + risk_premium

    st.subheader("Recommended Terms")

    formatted_loan = f"{adjusted_loan:,.0f}"
    st.success(f"Recommended Loan Amount: ₹ {formatted_loan}")

    # SAFE WORD CONVERSION (Overflow Protected)

    try:
        if adjusted_loan < 1e12:
            loan_in_words = num2words(int(adjusted_loan), lang='en_IN').title()
            st.info(f"Rupees {loan_in_words} Only")
        else:
            loan_in_crore = adjusted_loan / 1e7
            st.info(f"≈ ₹ {loan_in_crore:.2f} Crore")
    except:
        st.info("Loan amount too large to convert into words.")

    st.write("Recommended Interest Rate:", round(interest_rate, 2), "%")

    # ------------------------------------------------------
    # SHAP EXPLANATION
    # ------------------------------------------------------

    st.subheader("Explainability (Why this decision?)")

    shap_values = explainer(input_df)

    fig, ax = plt.subplots()
    shap.plots.waterfall(shap_values[0], show=False)

    st.pyplot(fig)

    # ------------------------------------------------------
    # CAM EXPORT (Word Document)
    # ------------------------------------------------------
    st.subheader("Generate Credit Appraisal Memo (CAM)")
    
    cam_doc = generate_cam_word(
        company_name=company_name,
        revenue=revenue,
        ebitda=ebitda,
        debt=debt,
        prob=prob,
        sentiment=sentiment_score,
        adjusted_loan=adjusted_loan,
        interest_rate=interest_rate,
        articles=articles
    )
    
    st.download_button(
        label="📄 Download 5-C Credit Appraisal Memo (Word)",
        data=cam_doc,
        file_name=f"{company_name.replace(' ', '_')}_CAM.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
